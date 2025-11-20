import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import RGBColor
import io
import datetime

# --- 核心格式控制函数 ---

def set_style(run, font_size=12, bold=False):
    """
    设置字体样式：
    - 中文：宋体
    - 西文：Times New Roman
    - 字号：默认小四 (12pt)
    """
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    run.font.bold = bold

def set_cell_text(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=12, bold=False, line_spacing=None):
    """
    设置表格单元格内容的通用函数
    """
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    
    # 设置行距
    if line_spacing:
        p.paragraph_format.line_spacing_rule = line_spacing
    else:
        # 表格内默认单倍行距，防止撑太大，除非特意指定
        p.paragraph_format.line_spacing = 1.2 

    run = p.add_run(text)
    set_style(run, font_size=font_size, bold=bold)
    return cell

def format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False):
    """设置段落基本格式：1.5倍行距"""
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74) # 首行缩进约2字符

# --- 页面生成逻辑 ---

def create_cover(doc, data):
    """
    生成封面 (Page 1)
    注：根据要求，Logo部分已简化，重点保证文字布局。
    """
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 学号 (左上角)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"★                 学号：{data['student_id']}")
    set_style(run, font_size=12) # 小四

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("岗前综合技能培训报告书")
    set_style(run, font_size=36, bold=True) # 一号/小初
    p.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()
    doc.add_paragraph()

    # 项目名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"项目： {data['project_name']} ")
    set_style(run, font_size=16, bold=True) # 三号
    run.font.underline = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 封面信息表 (无边框表格布局)
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 字段列表
    info_list = [
        ("学    院：", data['college']),
        ("专    业：", data['major']),
        ("班    级：", data['class_name']),
        ("学生姓名：", data['name']),
        ("指导教师：", data['teacher'])
    ]

    for row_idx, (label, val) in enumerate(info_list):
        # 标签列
        cell_label = table.cell(row_idx, 0)
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        set_style(run, font_size=16, bold=True) # 三号加粗
        
        # 内容列
        cell_val = table.cell(row_idx, 1)
        p = cell_val.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f" {val} ")
        set_style(run, font_size=16) # 三号
        run.font.underline = True
        
    doc.add_paragraph()
    doc.add_paragraph()

    # 时间
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"起止时间： {data['start_date']} 至 {data['end_date']}")
    set_style(run, font_size=14) # 四号

    doc.add_paragraph()
    doc.add_paragraph()
    
    # 底部学校名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("海南软件职业技术学院")
    set_style(run, font_size=22, bold=True) # 二号

    doc.add_page_break()

def create_task_sheet(doc, data):
    """生成任务书 (Page 2) - 重点优化"""
    
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 标题通常不需要1.5倍行距，单倍即可
    run = p.add_run("海南软件职业技术学院  岗前综合技能培训任务书")
    set_style(run, font_size=16, bold=True) # 三号

    # 创建表格：8行6列
    table = doc.add_table(rows=8, cols=6)
    table.style = 'Table Grid'
    table.autofit = False
    
    # 手动设置列宽以匹配视觉比例 (总宽约16-17cm)
    # 列宽：标题列略窄，内容列略宽
    col_widths = [Cm(2.5), Cm(3.5), Cm(2.0), Cm(2.5), Cm(2.0), Cm(3.5)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # --- 第一行：学院 |Val| 专业 |Val(合并3列)| ---
    set_cell_text(table.cell(0, 0), "学院")
    set_cell_text(table.cell(0, 1), data['college'])
    set_cell_text(table.cell(0, 2), "专业")
    table.cell(0, 3).merge(table.cell(0, 5))
    set_cell_text(table.cell(0, 3), data['major'])

    # --- 第二行：班级 |Val| 学号 |Val| 姓名 |Val| ---
    set_cell_text(table.cell(1, 0), "班级")
    set_cell_text(table.cell(1, 1), data['class_name'])
    set_cell_text(table.cell(1, 2), "学号")
    set_cell_text(table.cell(1, 3), data['student_id'])
    set_cell_text(table.cell(1, 4), "姓名")
    set_cell_text(table.cell(1, 5), data['name'])

    # --- 第三行：指导教师 |Val| 题目 |Val(合并3列)| ---
    set_cell_text(table.cell(2, 0), "岗前综合技能\n培训指导教师")
    set_cell_text(table.cell(2, 1), data['teacher'])
    set_cell_text(table.cell(2, 2), "题目")
    table.cell(2, 3).merge(table.cell(2, 5))
    set_cell_text(table.cell(2, 3), data['project_name'])

    # --- 第四行：起止时间 (合并后5列) ---
    set_cell_text(table.cell(3, 0), "岗前综合技能培训\n起止时间")
    table.cell(3, 1).merge(table.cell(3, 5))
    set_cell_text(table.cell(3, 1), f"{data['start_date']} 至 {data['end_date']}")

    # --- 大段文本区域 (小四宋体，靠左对齐) ---
    
    # 第五行：意义及目标
    set_cell_text(table.cell(4, 0), "项目的意义\n及培养目标")
    table.cell(4, 1).merge(table.cell(4, 5))
    set_cell_text(table.cell(4, 1), data['meaning'], align=WD_ALIGN_PARAGRAPH.LEFT)
    table.rows[4].height = Cm(2.5) # 设置最小行高

    # 第六行：成果形式
    set_cell_text(table.cell(5, 0), "岗前综合技能\n培训成果形式")
    table.cell(5, 1).merge(table.cell(5, 5))
    set_cell_text(table.cell(5, 1), data['output_form'], align=WD_ALIGN_PARAGRAPH.LEFT)
    table.rows[5].height = Cm(1.5)

    # 第七行：基本要求
    set_cell_text(table.cell(6, 0), "技能训练\n基本要求")
    table.cell(6, 1).merge(table.cell(6, 5))
    set_cell_text(table.cell(6, 1), data['requirements'], align=WD_ALIGN_PARAGRAPH.LEFT)
    table.rows[6].height = Cm(2.5)

    # 第八行：主要任务
    set_cell_text(table.cell(7, 0), "岗前综合技能\n培训主要任务")
    table.cell(7, 1).merge(table.cell(7, 5))
    set_cell_text(table.cell(7, 1), data['main_tasks'], align=WD_ALIGN_PARAGRAPH.LEFT)
    table.rows[7].height = Cm(3.0)

    # 底部签名区域
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("岗前综合技能培训指导教师签名：")
    set_style(run, font_size=12)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("岗前综合技能培训领导小组审查意见：")
    set_style(run, font_size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("组长签名：                             年    月    日")
    set_style(run, font_size=12)
    
    # 备注
    p = doc.add_paragraph("备注：此表回收后交院部按班级为单位装订存档。")
    set_style(p.runs[0], font_size=10.5) # 五号字

    doc.add_page_break()

def create_guidance_record(doc, data):
    """生成指导记录表 (Page 3)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("岗前综合技能培训指导记录表")
    set_style(run, font_size=16, bold=True)

    table = doc.add_table(rows=12, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    headers = [
        ("学    号", data['student_id'], "岗前综合技能\n培训指导教师", data['teacher']),
        ("姓    名", data['name'], "专    业", data['major']),
        ("班    级", data['class_name'], "指导教师专业\n方    向", ""), # 假设方向留空或手动填
    ]

    # 填充前三行基础信息
    for i, row_dat in enumerate(headers):
        set_cell_text(table.cell(i, 0), row_dat[0])
        set_cell_text(table.cell(i, 1), row_dat[1])
        set_cell_text(table.cell(i, 2), row_dat[2])
        set_cell_text(table.cell(i, 3), row_dat[3])

    # 第四行：项目名称（单独处理，可能需要合并）
    set_cell_text(table.cell(3, 0), "项目名称")
    table.cell(3, 1).merge(table.cell(3, 3))
    set_cell_text(table.cell(3, 1), data['project_name'])

    # 第五行：标题行
    set_cell_text(table.cell(4, 0), "指导时间")
    table.cell(4, 1).merge(table.cell(4, 3))
    set_cell_text(table.cell(4, 1), "指导内容")

    # 循环生成指导记录行 (7行)
    start_date = datetime.datetime.strptime(data['start_date'].split('至')[0].strip(), "%Y年%m月%d日") if '年' in data['start_date'] else datetime.datetime.now()
    
    for i in range(7):
        row_idx = 5 + i
        sim_date = start_date + datetime.timedelta(days=i*7)
        date_str = f"{sim_date.month}月{sim_date.day}日"
        
        set_cell_text(table.cell(row_idx, 0), date_str)
        table.cell(row_idx, 1).merge(table.cell(row_idx, 3))
        set_cell_text(table.cell(row_idx, 1), " ", align=WD_ALIGN_PARAGRAPH.LEFT) # 留空供手写
        table.rows[row_idx].height = Cm(1.2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("岗前综合技能培训指导教师（签名）：           年    月    日")
    set_style(run, font_size=12)

    doc.add_page_break()

def create_assessment(doc, data):
    """生成成绩评定表 (Page 4)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("海南软件职业技术学院\n岗前综合技能培训成绩评定表")
    set_style(run, font_size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"学院：{data['college']}   专业（班级）：{data['major']} {data['class_name']}   学号：{data['student_id']}   姓名：{data['name']}")
    set_style(run, font_size=10.5)

    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'

    # 标题行
    table.cell(0, 0).merge(table.cell(0, 4))
    set_cell_text(table.cell(0, 0), "岗前综合技能培训成绩评定", bold=True)

    # 项目名称
    set_cell_text(table.cell(1, 0), "岗前综合技能\n培训项目名称：")
    table.cell(1, 1).merge(table.cell(1, 4))
    set_cell_text(table.cell(1, 1), data['project_name'])

    # 成果
    set_cell_text(table.cell(2, 0), "岗前综合技能\n培训成果：")
    table.cell(2, 1).merge(table.cell(2, 4))
    set_cell_text(table.cell(2, 1), f"□软件作品  □影视动漫作品  □电子工艺产品  □综述报告\n□项目文档  □文艺作品      □其他：{data['output_form']}", align=WD_ALIGN_PARAGRAPH.LEFT)

    # 指导教师评语
    set_cell_text(table.cell(3, 0), "岗前综合技能\n培训指导教师\n评语")
    table.cell(3, 1).merge(table.cell(3, 4))
    set_cell_text(table.cell(3, 1), "\n\n\n\n指导教师签名：          ", align=WD_ALIGN_PARAGRAPH.RIGHT)
    table.rows[3].height = Cm(3)

    # 初评成绩
    set_cell_text(table.cell(4, 0), "审查时间")
    set_cell_text(table.cell(4, 1), "    年  月  日")
    set_cell_text(table.cell(4, 2), "初评成绩")
    table.cell(4, 3).merge(table.cell(4, 4))
    set_cell_text(table.cell(4, 3), " ")

    # 答辩 (结构较复杂，简化处理以匹配)
    table.cell(5, 0).merge(table.cell(6, 0))
    set_cell_text(table.cell(5, 0), "岗前综合技能\n培训优秀项目\n（答辩）\n成绩")
    
    # 评分标准1
    table.cell(5, 1).merge(table.cell(5, 3))
    set_cell_text(table.cell(5, 1), "1．岗前综合技能培训成果水平和工作量评价（满分 80 分）\nA. 有创新性结果，全面完成了训练任务 (71-80 分)\nB. 有创新性结果，基本完成了训练任务 (61-70 分)\nC. 有一定的创新性结果，基本完成了训练任务 (51-60 分)\nD. 基本没有创新性结果，没有完成训练任务 (0-50 分)", align=WD_ALIGN_PARAGRAPH.LEFT, font_size=9)
    set_cell_text(table.cell(5, 4), "评分")

    # 评分标准2
    table.cell(6, 1).merge(table.cell(6, 3))
    set_cell_text(table.cell(6, 1), "2．答辩材料准备与答辩表现（满分 20 分）\nA. 准备充分，展示较好，概念清楚 (15-20 分)\nB. 展示及回答问题表现较好 (10-15 分)\nC. 展示及回答问题表现一般 (5-10 分)\nD. 展示及回答问题表现很差 (0-5 分)", align=WD_ALIGN_PARAGRAPH.LEFT, font_size=9)
    set_cell_text(table.cell(6, 4), "评分")

    # 评语/签名
    set_cell_text(table.cell(7, 0), "评语")
    set_cell_text(table.cell(7, 1), " ")
    set_cell_text(table.cell(7, 2), "答辩成绩")
    set_cell_text(table.cell(7, 3), " ")
    set_cell_text(table.cell(7, 4), "答辩小组负责人签名：\n\n    年  月  日")

    # 最终成绩
    set_cell_text(table.cell(8, 0), "岗前综合技能\n培训\n最终成绩评定")
    table.cell(8, 1).merge(table.cell(8, 4))
    set_cell_text(table.cell(8, 1), "成绩评定（在“□”中划“ √”)\n\n优□    良□    中□    及格□    不及格□\n\n学院（签章）：                     年    月    日", align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

def create_approval_form(doc, data):
    """生成选题审批表 (Page 5)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("海南软件职业技术学院\n岗前综合技能培训选题审批表")
    set_style(run, font_size=16, bold=True)

    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'

    # 第一行
    set_cell_text(table.cell(0, 0), "学  号")
    set_cell_text(table.cell(0, 1), data['student_id'])
    set_cell_text(table.cell(0, 2), "姓   名")
    set_cell_text(table.cell(0, 3), data['name'])
    set_cell_text(table.cell(0, 4), "班  级")
    set_cell_text(table.cell(0, 5), data['class_name'])

    # 第二行
    set_cell_text(table.cell(1, 0), "岗前综合技能\n培训项目名称")
    table.cell(1, 1).merge(table.cell(1, 3))
    set_cell_text(table.cell(1, 1), data['project_name'])
    set_cell_text(table.cell(1, 4), "指导教师")
    set_cell_text(table.cell(1, 5), data['teacher'])

    # 第三行：选题理由
    set_cell_text(table.cell(2, 0), "选题理由及\n准备情况：")
    table.cell(2, 1).merge(table.cell(2, 5))
    set_cell_text(table.cell(2, 1), data['reason'], align=WD_ALIGN_PARAGRAPH.LEFT)
    table.rows[2].height = Cm(5)

    # 第四行：指导教师意见
    set_cell_text(table.cell(3, 0), "指导教师意见：")
    table.cell(3, 1).merge(table.cell(3, 5))
    set_cell_text(table.cell(3, 1), "\n\n指导教师（签字）           年      月     日", align=WD_ALIGN_PARAGRAPH.RIGHT)
    table.rows[3].height = Cm(3)

    # 第五行：学院意见
    set_cell_text(table.cell(4, 0), "学院意见：")
    table.cell(4, 1).merge(table.cell(4, 5))
    set_cell_text(table.cell(4, 1), "\n\n学院（签字）              年      月     日", align=WD_ALIGN_PARAGRAPH.RIGHT)
    table.rows[4].height = Cm(3)

    doc.add_page_break()

def create_report_body_template(doc, data):
    """生成报告正文模板 (Page 6+) - 1.5倍行距，宋体小四"""
    
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{data['project_name']} 培训报告")
    set_style(run, font_size=16, bold=True) # 三号

    doc.add_paragraph()

    # 一、目的
    p = doc.add_paragraph()
    run = p.add_run("一、岗前培训目的")
    set_style(run, font_size=14, bold=True) # 四号加粗
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

    p = doc.add_paragraph("（在此处撰写岗前培训的目的和意义，岗前培训单位的发展情况及学习要求等，不少于300字。）")
    set_style(p.runs[0], font_size=12) # 小四
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    doc.add_paragraph() 

    # 二、内容
    p = doc.add_paragraph()
    run = p.add_run("二、岗前培训内容")
    set_style(run, font_size=14, bold=True)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = doc.add_paragraph("1、项目背景与需求分析")
    set_style(p.runs[0], font_size=12, bold=True) # 小四加粗
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = doc.add_paragraph("（正文内容：请详细描述项目的具体内容、所用技术栈、开发过程中的关键步骤等。注意字体为宋体小四，行距为1.5倍。）")
    set_style(p.runs[0], font_size=12)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = doc.add_paragraph("2、系统设计与实现")
    set_style(p.runs[0], font_size=12, bold=True)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    p = doc.add_paragraph("（正文内容...）")
    set_style(p.runs[0], font_size=12)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    doc.add_paragraph()

    # 三、结果
    p = doc.add_paragraph()
    run = p.add_run("三、岗前培训结果")
    set_style(run, font_size=14, bold=True)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    p = doc.add_paragraph("（展示运行结果、界面截图等。）")
    set_style(p.runs[0], font_size=12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    doc.add_paragraph()

    # 四、总结
    p = doc.add_paragraph()
    run = p.add_run("四、培训总结或体会")
    set_style(run, font_size=14, bold=True)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    p = doc.add_paragraph("（总结培训过程中的收获、遇到的问题及解决方案、不足以及对未来职业生涯的规划。）")
    set_style(p.runs[0], font_size=12)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# --- Streamlit UI ---

st.set_page_config(page_title="岗前培训报告生成工具", layout="wide")
st.title("📝 岗前综合技能培训报告生成工具")
st.markdown("""
**适用院校：** 海南软件职业技术学院  
**格式标准：** 宋体小四，1.5倍行距，标准表格样式  
**当前状态：** 已优化表格结构，匹配官方模板
""")

with st.sidebar:
    st.header("1. 基本信息输入")
    name = st.text_input("学生姓名", "张三")
    student_id = st.text_input("学号", "20230001")
    college = st.text_input("学院", "机电工程学院")
    major = st.text_input("专业", "软件技术")
    class_name = st.text_input("班级", "23软件技术1班")
    teacher = st.text_input("指导教师", "李四")
    
    st.header("2. 项目信息")
    project_name = st.text_input("项目名称", "基于Python的企业网站开发")
    today = datetime.date.today()
    start_date_obj = st.date_input("开始时间", datetime.date(2025, 7, 1))
    end_date_obj = st.date_input("结束时间", datetime.date(2025, 8, 31))
    start_date = start_date_obj.strftime("%Y年%m月%d日")
    end_date = end_date_obj.strftime("%Y年%m月%d日")
    
    st.header("3. 任务书详细内容")
    meaning = st.text_area("项目的意义及培养目标", "通过本项目训练，掌握Web开发全流程，提升编码能力...", height=100)
    output_form = st.selectbox("成果形式", ["软件作品", "项目文档", "综述报告", "电子工艺产品", "其他"])
    requirements = st.text_area("技能训练基本要求", "1. 代码规范\n2. 功能完整\n3. 文档齐全", height=100)
    main_tasks = st.text_area("主要任务", "1. 需求分析\n2. 数据库设计\n3. 前端页面开发\n4. 后端接口实现", height=100)
    
    st.header("4. 审批表内容")
    reason = st.text_area("选题理由", "该项目符合专业培养目标，且能结合实习岗位实际...", height=80)

# 数据打包
data = {
    "name": name, "student_id": student_id, "college": college,
    "major": major, "class_name": class_name, "teacher": teacher,
    "project_name": project_name, "start_date": start_date, "end_date": end_date,
    "meaning": meaning, "output_form": output_form, 
    "requirements": requirements, "main_tasks": main_tasks,
    "reason": reason
}

# 主界面
st.info("👈 请在左侧侧边栏完善信息。")

col1, col2 = st.columns([2, 1])
with col1:
    st.subheader("包含的模板页面")
    st.markdown("""
    - **封面**：标准格式（暂无LOGO）
    - **任务书**：包含意义、目标、任务等大段文本，表格布局已优化
    - **指导记录表**：生成7-8周的记录模板
    - **成绩评定表**：包含答辩评分细则
    - **选题审批表**：包含选题理由和意见栏
    - **正文模板**：**宋体小四 + 1.5倍行距**
    """)

with col2:
    st.write("### 操作")
    if st.button("🚀 生成标准Word报告", type="primary"):
        doc = Document()
        
        # 全局设置默认字体（备用）
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(12)
        
        create_cover(doc, data)
        create_task_sheet(doc, data)
        create_guidance_record(doc, data)
        create_assessment(doc, data)
        create_approval_form(doc, data)
        create_report_body_template(doc, data)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        file_name = f"{student_id}_{name}_岗前培训报告.docx"
        st.success(f"文档已生成！文件名为：{file_name}")
        st.download_button(
            label="📥 下载 .docx 文件",
            data=buffer,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
